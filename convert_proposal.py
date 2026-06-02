import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4


def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    doc = Document()
    for line in lines:
        text = line.rstrip('\n')
        if text.startswith('# '):
            doc.add_heading(text[2:].strip(), level=1)
        elif text.startswith('## '):
            doc.add_heading(text[3:].strip(), level=2)
        elif text.startswith('- '):
            doc.add_paragraph(text[2:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(text)
    doc.save(docx_path)


def md_to_pdf(md_path, pdf_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    margin = 40
    y = height - margin
    max_width = width - 2 * margin
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    try:
        pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
        fontname = 'DejaVuSans'
    except Exception:
        fontname = 'Helvetica'
    line_height = 14
    for raw in lines:
        text = raw.rstrip('\n')
        if text.startswith('# '):
            txt = text[2:].strip()
            size = 16
            c.setFont(fontname, size)
        elif text.startswith('## '):
            txt = text[3:].strip()
            size = 14
            c.setFont(fontname, size)
        else:
            txt = text
            size = 11
            c.setFont(fontname, size)
        from reportlab.pdfbase.pdfmetrics import stringWidth
        parts = []
        words = txt.split(' ')
        cur = ''
        for w in words:
            test = (cur + ' ' + w).strip()
            if stringWidth(test, fontname, size) <= max_width:
                cur = test
            else:
                parts.append(cur)
                cur = w
        if cur:
            parts.append(cur)
        for p in parts:
            if y < margin + line_height:
                c.showPage()
                y = height - margin
                c.setFont(fontname, size)
            c.drawString(margin, y, p)
            y -= line_height
        y -= line_height / 2
    c.save()


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    md = os.path.join(base, 'proposta_professor.md')
    docx = os.path.join(base, 'proposta_professor.docx')
    pdf = os.path.join(base, 'proposta_professor.pdf')
    md_to_docx(md, docx)
    md_to_pdf(md, pdf)
    print('Wrote', docx)
    print('Wrote', pdf)
